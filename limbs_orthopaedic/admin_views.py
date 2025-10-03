from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout
from django.contrib.auth.models import User
from django.views.decorators.http import require_POST
from django.db.models import Count, Q, Sum
from django.db.models.functions import TruncMonth, TruncWeek, TruncDay
from django.utils import timezone
from django.http import JsonResponse, HttpResponse
from django.contrib import messages
from datetime import timedelta, datetime
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

from accounts.models import Doctor, Profile
from analytics.models import PageVisit
from appointments.models import Appointment
from blog.models import BlogPost, BlogCategory, Comment
from products.models import CustomProductRequest, Product
from core.models import Contact
from surveys.models import Survey, Response, Question, Answer
from services.models import Service
from gallery.models import Gallery 
from staff_chat.models import StaffChatMessage
from testimonials.models import Testimonial
from careers.models import Career, CareerApplication
from limbs_orthopaedic.models import Invoice, InvoiceItem
from django.utils import timezone
from datetime import datetime

@staff_member_required
def admin_dashboard(request):
    """Admin dashboard view with statistics and recent data"""

    # Get time filter from query params
    time_filter = request.GET.get('time_filter', 'all')

    # Define date range based on filter
    now = timezone.now()
    if time_filter == 'year':
        start_date = now - timedelta(days=365)
    elif time_filter == 'month':
        start_date = now - timedelta(days=30)
    elif time_filter == 'week':
        start_date = now - timedelta(days=7)
    elif time_filter == 'day':
        start_date = now - timedelta(days=1)
    else:
        start_date = None  # All time

    # Base querysets - can be filtered by date later
    users_qs = User.objects.all()
    appointments_qs = Appointment.objects.all()
    surveys_qs = Response.objects.filter(is_complete=True)

    # Apply date filters if needed
    if start_date:
        users_qs = users_qs.filter(date_joined__gte=start_date)
        appointments_qs = appointments_qs.filter(preferred_date__gte=start_date)
        surveys_qs = surveys_qs.filter(completed_at__gte=start_date)

    # Get staff messages count

    # Get counts for statistics
    total_users = User.objects.count()  # All registered users
    doctor_count = Doctor.objects.count()  # Count all doctors
    appointment_count = appointments_qs.count()
    blog_count = BlogPost.objects.count()
    blog_comments_count = Comment.objects.count()  # Count blog comments
    survey_count = Survey.objects.count()  # Count all surveys
    surveys_done_count = surveys_qs.count()  # Count completed surveys
    service_count = Service.objects.count()  # Count services
    contact_count = Contact.objects.count()  # Count contact requests
    testimonial_count = Testimonial.objects.count()  # Count testimonials
    
    # Get career statistics
    careers = Career.objects.filter(is_active=True)
    now = timezone.now()
    for career in careers:
        career.applications_this_month = CareerApplication.objects.filter(
            career=career,
            created_at__year=now.year,
            created_at__month=now.month
        ).count()
        career.applications_this_year = CareerApplication.objects.filter(
            career=career,
            created_at__year=now.year
        ).count()
        career.total_applications = CareerApplication.objects.filter(career=career).count()
    product_count = Product.objects.filter(is_active=True).count()  # Count active products
    product_inquiry_count = CustomProductRequest.objects.count()  # Count product inquiries
    gallery_count = Gallery.objects.filter(is_active=True).count()  # Count active gallery items
    male_users = Profile.objects.filter(gender='M').count()
    female_users = Profile.objects.filter(gender='F').count()
    gender_users = male_users + female_users  # Users who have set their gender
    staff_messages_count = StaffChatMessage.objects.count()  # Count staff messages


    # Get recent appointments
    recent_appointments = appointments_qs.order_by('-preferred_date')[:5]

    # Get doctors
    doctors = Doctor.objects.filter(is_active=True)[:5]

    # Get recent blog posts
    recent_blog_posts = BlogPost.objects.all().order_by('-published_date')[:5]

    # Get recent custom product requests
    recent_custom_product_requests = CustomProductRequest.objects.all().order_by('-created_at')[:5]

    # Get recent contact messages
    recent_contact_messages = Contact.objects.all().order_by('-created_at')[:5]

    # Get recent survey responses
    recent_survey_responses = Response.objects.filter(is_complete=True).order_by('-completed_at')[:5]

    # Prepare data for visualizations
    # 1. Appointments by month
    appointments_by_month = []

    # Get months data
    month_data = appointments_qs.annotate(
        month=TruncMonth('preferred_date')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            appointments_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 2. Survey responses by month
    surveys_by_month = []

    # Get survey response data by month
    survey_month_data = surveys_qs.annotate(
        month=TruncMonth('completed_at')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in survey_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            surveys_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 3. Custom Products by month
    # Using the already imported CustomProductRequest from line 15
    custom_products_qs = CustomProductRequest.objects.all()
    if start_date:
        custom_products_qs = custom_products_qs.filter(created_at__gte=start_date)

    custom_products_by_month = []
    custom_products_month_data = custom_products_qs.annotate(
        month=TruncMonth('created_at')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in custom_products_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            custom_products_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 4. Testimonials by month
    testimonials_qs = Testimonial.objects.all()
    if start_date:
        testimonials_qs = testimonials_qs.filter(created_at__gte=start_date)

    testimonials_by_month = []
    testimonials_month_data = testimonials_qs.annotate(
        month=TruncMonth('created_at')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in testimonials_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            testimonials_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 5. Contact messages by month
    # Using the already imported Contact from line 16
    contacts_qs = Contact.objects.all()
    if start_date:
        contacts_qs = contacts_qs.filter(created_at__gte=start_date)

    contacts_by_month = []
    contacts_month_data = contacts_qs.annotate(
        month=TruncMonth('created_at')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in contacts_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            contacts_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 6. User registrations by month
    # Using the already imported User model from line 4
    users_reg_qs = User.objects.all()
    if start_date:
        users_reg_qs = users_reg_qs.filter(date_joined__gte=start_date)

    users_by_month = []
    users_month_data = users_reg_qs.annotate(
        month=TruncMonth('date_joined')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in users_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            users_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 7. Blog posts by month
    # Using the already imported BlogPost and Comment models from line 14
    blogs_qs = BlogPost.objects.all()
    if start_date:
        blogs_qs = blogs_qs.filter(published_date__gte=start_date)

    blogs_by_month = []
    blogs_month_data = blogs_qs.annotate(
        month=TruncMonth('published_date')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in blogs_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            blogs_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # 8. Blog comments by month
    comments_qs = Comment.objects.all()
    if start_date:
        comments_qs = comments_qs.filter(created_date__gte=start_date)

    comments_by_month = []
    comments_month_data = comments_qs.annotate(
        month=TruncMonth('created_date')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')

    for data in comments_month_data:
        if data['month']:
            month_name = data['month'].strftime('%b %Y')
            comments_by_month.append({
                'month': month_name,
                'count': data['count']
            })

    # Prepare chart data as JSON
    chart_data = {
        'appointments_by_month': json.dumps(appointments_by_month),
        'surveys_by_month': json.dumps(surveys_by_month),
        'custom_products_by_month': json.dumps(custom_products_by_month),
        'testimonials_by_month': json.dumps(testimonials_by_month),
        'contacts_by_month': json.dumps(contacts_by_month),
        'users_by_month': json.dumps(users_by_month),
        'blogs_by_month': json.dumps(blogs_by_month),
        'comments_by_month': json.dumps(comments_by_month),
    }

    # Get visit statistics
    visit_counts = PageVisit.get_visit_counts()
    
    context = {
        'visit_counts': visit_counts,
        'total_users': total_users,
        'male_users': male_users,
        'female_users': female_users,
        'doctor_count': doctor_count,
        'appointment_count': appointment_count,
        'blog_count': blog_count,
        'blog_comments_count': blog_comments_count,
        'survey_count': survey_count,
        'surveys_done_count': surveys_done_count,
        'service_count': service_count,
        'contact_count': contact_count,
        'testimonial_count': testimonial_count,
        'product_count': product_count,
        'product_inquiry_count': product_inquiry_count,
        'gallery_count': gallery_count,
        'staff_messages_count': staff_messages_count,
        'recent_appointments': recent_appointments,
        'doctors': doctors,
        'careers': careers,
        'recent_blog_posts': recent_blog_posts,
        'recent_custom_product_requests': recent_custom_product_requests,
        'recent_contact_messages': recent_contact_messages,
        'recent_survey_responses': recent_survey_responses,
        'chart_data': chart_data,
        'time_filter': time_filter,
    }

    return render(request, 'admin/dashboard.html', context)


@staff_member_required
def admin_invoice_generator(request):
    """Invoice generator view for admin users"""
    try:
        # Get actual statistics from database
        invoices = Invoice.objects.all()
        total_invoices = invoices.count()
        total_revenue = invoices.aggregate(total=Sum('total_amount'))['total'] or 0
        last_invoice = invoices.first()
        
        context = {
            'title': 'Invoice Generator',
            'site_title': 'LIMBS Orthopaedic Admin',
            'site_header': 'LIMBS Orthopaedic Administration',
            'total_invoices': total_invoices,
            'total_revenue': total_revenue,
            'last_invoice': last_invoice,
        }
        return render(request, 'admin/invoice_generator.html', context)
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error in invoice generator view: {str(e)}")
        from django.http import HttpResponse
        return HttpResponse(f"Error loading invoice generator: {str(e)}", status=500)


@staff_member_required
def admin_invoice_list(request):
    """List all invoices"""
    invoices = Invoice.objects.all().order_by('-created_at')
    
    # Handle search
    search_query = request.GET.get('search', '')
    if search_query:
        invoices = invoices.filter(
            Q(invoice_number__icontains=search_query) |
            Q(patient_name__icontains=search_query) |
            Q(tracking_code__icontains=search_query)
        )
    
    # Handle status filter
    status_filter = request.GET.get('status', '')
    if status_filter:
        invoices = invoices.filter(status=status_filter)
    
    context = {
        'title': 'Invoice List',
        'invoices': invoices,
        'search_query': search_query,
        'status_filter': status_filter,
        'status_choices': Invoice.STATUS_CHOICES,
    }
    return render(request, 'admin/invoice_list.html', context)


@staff_member_required
def admin_invoice_detail(request, invoice_id):
    """View invoice details"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    context = {
        'title': f'Invoice {invoice.invoice_number}',
        'invoice': invoice,
    }
    return render(request, 'admin/invoice_detail.html', context)


@staff_member_required
def admin_invoice_edit(request, invoice_id):
    """Edit invoice"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    
    if request.method == 'POST':
        try:
            # Update invoice data
            invoice.patient_name = request.POST.get('patient_name')
            invoice.patient_address = request.POST.get('patient_address')
            invoice.patient_phone = request.POST.get('patient_phone')
            invoice.status = request.POST.get('status', 'draft')
            invoice.notes = request.POST.get('notes', '')
            
            # Clear existing items and add new ones
            invoice.items.all().delete()
            
            subtotal = 0
            services_data = json.loads(request.POST.get('services_data', '[]'))
            
            for service in services_data:
                if service.get('name') and service.get('unitPrice', 0) > 0:
                    quantity = int(service.get('quantity', 1))
                    unit_price = float(service.get('unitPrice', 0))
                    
                    InvoiceItem.objects.create(
                        invoice=invoice,
                        description=service['name'],
                        quantity=quantity,
                        unit_price=unit_price,
                        total_price=quantity * unit_price
                    )
                    subtotal += quantity * unit_price
            
            # Update totals
            tax_amount = subtotal * 0.16  # 16% VAT
            invoice.subtotal = subtotal
            invoice.tax_amount = tax_amount
            invoice.total_amount = subtotal + tax_amount
            invoice.save()
            
            messages.success(request, f'Invoice {invoice.invoice_number} updated successfully!')
            return redirect('admin_invoice_detail', invoice_id=invoice.id)
            
        except Exception as e:
            messages.error(request, f'Error updating invoice: {str(e)}')
    
    # Prepare services data for the form
    services_data = []
    for item in invoice.items.all():
        services_data.append({
            'name': item.description,
            'quantity': item.quantity,
            'unitPrice': float(item.unit_price),
            'totalPrice': float(item.total_price)
        })
    
    context = {
        'title': f'Edit Invoice {invoice.invoice_number}',
        'invoice': invoice,
        'services_data': json.dumps(services_data),
    }
    return render(request, 'admin/invoice_edit.html', context)


@staff_member_required
@require_POST
def admin_save_invoice(request):
    """Save a new invoice to database"""
    try:
        # Get form data
        patient_name = request.POST.get('patient_name', '').strip()
        patient_address = request.POST.get('patient_address', '').strip()
        patient_phone = request.POST.get('patient_phone', '').strip()
        
        # Validate required fields
        if not all([patient_name, patient_address, patient_phone]):
            return JsonResponse({'success': False, 'error': 'All patient fields are required'})
        
        # Parse services data
        services_data = json.loads(request.POST.get('services_data', '[]'))
        if not services_data:
            return JsonResponse({'success': False, 'error': 'At least one service is required'})
        
        # Generate invoice number and tracking code
        now = datetime.now()
        date_str = now.strftime('%Y%m%d')
        today_invoices = Invoice.objects.filter(created_at__date=now.date()).count()
        sequence = str(today_invoices + 1).zfill(3)
        invoice_number = f"LOL-{date_str}-{sequence}"
        
        # Generate tracking code
        import uuid
        tracking_code = f"LOL-{uuid.uuid4().hex[:5].upper()}-{uuid.uuid4().hex[:5].upper()}"
        
        # Calculate totals
        subtotal = 0
        for service in services_data:
            if service.get('name') and service.get('unitPrice', 0) > 0:
                quantity = int(service.get('quantity', 1))
                unit_price = float(service.get('unitPrice', 0))
                subtotal += quantity * unit_price
        
        tax_amount = subtotal * 0.16  # 16% VAT
        total_amount = subtotal + tax_amount
        
        # Create invoice
        invoice = Invoice.objects.create(
            invoice_number=invoice_number,
            tracking_code=tracking_code,
            patient_name=patient_name,
            patient_address=patient_address,
            patient_phone=patient_phone,
            subtotal=subtotal,
            tax_amount=tax_amount,
            total_amount=total_amount,
            created_by=request.user
        )
        
        # Create invoice items
        for service in services_data:
            if service.get('name') and service.get('unitPrice', 0) > 0:
                quantity = int(service.get('quantity', 1))
                unit_price = float(service.get('unitPrice', 0))
                
                InvoiceItem.objects.create(
                    invoice=invoice,
                    description=service['name'],
                    quantity=quantity,
                    unit_price=unit_price,
                    total_price=quantity * unit_price
                )
        
        return JsonResponse({
            'success': True,
            'invoice_id': invoice.id,
            'invoice_number': invoice_number,
            'tracking_code': tracking_code,
            'total_amount': float(total_amount)
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@staff_member_required
def admin_invoice_delete(request, invoice_id):
    """Delete an invoice"""
    if request.method == 'POST':
        invoice = get_object_or_404(Invoice, id=invoice_id)
        invoice_number = invoice.invoice_number
        invoice.delete()
        messages.success(request, f'Invoice {invoice_number} deleted successfully!')
        return redirect('admin_invoice_list')
    
    return redirect('admin_invoice_list')


@staff_member_required
def admin_invoice_download(request, invoice_id):
    """Download invoice as Word document"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    
    # Create a new Word document
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add header with clinic info
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Left column - Clinic info
    left_cell = header_table.rows[0].cells[0]
    clinic_name = left_cell.paragraphs[0]
    clinic_name.text = "LIMBS ORTHOPAEDIC LIMITED"
    clinic_name.runs[0].font.size = Pt(24)
    clinic_name.runs[0].font.bold = True
    clinic_name.runs[0].font.color.rgb = RGBColor(52, 189, 242)
    
    left_cell.add_paragraph("ICIPE Road, Kasarani, Nairobi, Kenya")
    left_cell.add_paragraph("Phone: +254 705 347 657 | Email: info@limbsorthopaedic.org")
    left_cell.add_paragraph("Website: https://limbsorthopaedic.org")
    
    # Right column - Invoice details
    right_cell = header_table.rows[0].cells[1]
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    invoice_title = right_cell.paragraphs[0]
    invoice_title.text = "INVOICE"
    invoice_title.runs[0].font.size = Pt(24)
    invoice_title.runs[0].font.bold = True
    invoice_title.runs[0].font.color.rgb = RGBColor(52, 189, 242)
    
    inv_num = right_cell.add_paragraph(f"Invoice #: {invoice.invoice_number}")
    inv_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    inv_num.runs[0].font.bold = True
    
    inv_date = right_cell.add_paragraph(f"Date: {invoice.created_at.strftime('%b %d, %Y')}")
    inv_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    tracking = right_cell.add_paragraph(f"Tracking: {invoice.tracking_code}")
    tracking.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Patient information section
    patient_heading = doc.add_heading("BILL TO:", level=2)
    patient_heading.runs[0].font.color.rgb = RGBColor(52, 189, 242)
    
    doc.add_paragraph(invoice.patient_name).runs[0].font.bold = True
    doc.add_paragraph(invoice.patient_address)
    doc.add_paragraph(f"Phone: {invoice.patient_phone}")
    
    doc.add_paragraph()  # Spacing
    
    # Services table
    services_table = doc.add_table(rows=1, cols=4)
    services_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = services_table.rows[0].cells
    header_cells[0].text = 'Description'
    header_cells[1].text = 'Quantity'
    header_cells[2].text = 'Unit Price'
    header_cells[3].text = 'Total'
    
    # Make header bold
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Set cell background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '34BDF2')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add invoice items
    for item in invoice.items.all():
        row_cells = services_table.add_row().cells
        row_cells[0].text = item.description
        row_cells[1].text = str(item.quantity)
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].text = f"KSh {item.unit_price:,.2f}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[3].text = f"KSh {item.total_price:,.2f}"
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Totals section
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Subtotal
    totals_table.rows[0].cells[0].text = 'Subtotal:'
    totals_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    totals_table.rows[0].cells[1].text = f"KSh {invoice.subtotal:,.2f}"
    totals_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # VAT
    totals_table.rows[1].cells[0].text = 'VAT (16%):'
    totals_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    totals_table.rows[1].cells[1].text = f"KSh {invoice.tax_amount:,.2f}"
    totals_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Total
    totals_table.rows[2].cells[0].text = 'TOTAL:'
    totals_table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    totals_table.rows[2].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    totals_table.rows[2].cells[1].text = f"KSh {invoice.total_amount:,.2f}"
    totals_table.rows[2].cells[1].paragraphs[0].runs[0].font.bold = True
    totals_table.rows[2].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    totals_table.rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Green background for total row
    for cell in totals_table.rows[2].cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '27AE60')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()  # Spacing
    
    # Payment methods section
    payment_heading = doc.add_heading("PAYMENT METHODS", level=2)
    payment_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    payment_heading.runs[0].font.color.rgb = RGBColor(52, 189, 242)
    
    # Bank payment details
    bank_heading = doc.add_heading("1. BANK PAYMENT", level=3)
    bank_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph("REF: BANK DETAILS WITH REFERENCE TO THE ABOVE").runs[0].font.bold = True
    doc.add_paragraph("THE FOLLOWING ARE DETAILS FOR TRANSFER OF FUNDS IN KES.")
    
    doc.add_paragraph("1. BANK ACCOUNT NAME: LIMBS ORTHOPAEDIC LIMITED").runs[0].font.bold = True
    doc.add_paragraph("2. BANK ACCOUNT: 1290426139").runs[0].font.bold = True
    doc.add_paragraph("3. SWIFT CODE: KCBLKENX").runs[0].font.bold = True
    doc.add_paragraph("4. BANK CODE/BRANCH CODE: 01107").runs[0].font.bold = True
    doc.add_paragraph("5. BRANCH NAME: TOM MBOYA").runs[0].font.bold = True
    doc.add_paragraph("6. PIN: P052046452B").runs[0].font.bold = True
    
    # M-PESA details
    mpesa_heading = doc.add_heading("2. M-PESA PAYBILL", level=3)
    mpesa_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph("a). Paybill Number: 522533").runs[0].font.bold = True
    doc.add_paragraph("b). Account Number: 5807001").runs[0].font.bold = True
    doc.add_paragraph("c). Business Name: LIMBS ORTHOPAEDIC LTD").runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Footer
    footer = doc.add_paragraph("Thank you for choosing LIMBS Orthopaedic Limited for your orthopaedic care needs.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True
    
    contact_footer = doc.add_paragraph("For any questions about this invoice, please contact us at the above phone number or email address.")
    contact_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_footer.runs[0].font.size = Pt(10)
    
    # Save document to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=Invoice_{invoice.invoice_number}.docx'
    
    return response


@require_POST
def admin_logout_view(request):
    """Custom admin logout view that redirects to homepage"""
    logout(request)
    return redirect('/')